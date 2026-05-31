"""Gera o relatório final em Word (.docx) a partir do conteúdo revisado.
Espelha relatorio_final.tex. Rode com o python do .venv do projeto."""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = Path(__file__).parent / "relatorio_final.docx"

doc = Document()

# ---- estilos base ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

AZUL = RGBColor(0x1F, 0x3B, 0x57)


def add_heading(txt, level=1):
    h = doc.add_heading(txt, level=level)
    for run in h.runs:
        run.font.color.rgb = AZUL
    return h


def add_para(txt, bold=False, italic=False, align=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(txt)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    return p


def add_rich(p, segments):
    """segments: lista de (texto, bold, italic)."""
    for txt, b, i in segments:
        r = p.add_run(txt)
        r.bold = b
        r.italic = i


def add_bullet(segments):
    p = doc.add_paragraph(style="List Bullet")
    add_rich(p, segments)
    return p


def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


# =====================================================================
# CAPA
# =====================================================================
title = doc.add_heading("Mortalidade por Causas Evitáveis no Brasil", level=0)
for run in title.runs:
    run.font.color.rgb = AZUL
add_para("Desigualdade socioeconômica, padrões municipais e perfis de causa (2022–2024)",
         italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
add_para("Projeto Data Major — Tópicos em Banco de Dados — IESB",
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Maio de 2026", align=WD_ALIGN_PARAGRAPH.CENTER)

# ---- Resumo ----
add_heading("Resumo", level=1)
add_para(
    "Este trabalho investiga a relação entre a mortalidade por causas evitáveis e as condições "
    "socioeconômicas dos municípios brasileiros. A partir de um pipeline de dados completo (ETL), "
    "integramos os microdados de óbitos do SIM/DATASUS (2022–2024) a indicadores de desenvolvimento "
    "humano, renda, desigualdade e investimento em saúde. Aplicamos a Lista Brasileira de Causas "
    "Evitáveis (LBCE) para identificar aproximadamente 1,37 milhão de óbitos evitáveis na faixa de "
    "5 a 74 anos. Sobre esse conjunto realizamos três análises complementares: (i) regressão "
    "multivariada (OLS, Random Forest e XGBoost) para medir a associação entre variáveis "
    "socioeconômicas e a taxa de mortalidade evitável; (ii) clusterização (K-Means) dos municípios "
    "em quatro perfis socioeconômicos; e (iii) mineração das causas básicas (CID-10) para "
    "caracterizar de que morre cada perfil. Para garantir comparabilidade entre municípios com "
    "estruturas etárias distintas, calculamos a taxa de mortalidade padronizada por idade "
    "(padronização direta com a população do Censo 2022), adotada como desfecho principal. Os "
    "resultados confirmam associação estatisticamente significativa entre desigualdade (Gini) e "
    "mortalidade evitável, e revelam uma assinatura socioeconômica nas causas: municípios mais "
    "pobres concentram doenças infecciosas e negligenciadas, enquanto os mais desenvolvidos "
    "concentram doenças crônicas. A padronização por idade corrige um paradoxo das taxas brutas "
    "(Sul/Sudeste apareciam com mortalidade maior por terem população mais envelhecida) e alinha "
    "os sinais da regressão à hipótese socioeconômica.",
    italic=True,
)

# =====================================================================
# 1. INTRODUÇÃO
# =====================================================================
add_heading("1. Introdução", level=1)

p = doc.add_paragraph()
add_rich(p, [("Problema. ", True, False),
             ("Causas evitáveis de morte são aquelas que poderiam ser prevenidas ou tratadas por "
              "ações do Sistema Único de Saúde (SUS): imunização, diagnóstico precoce, atenção básica "
              "e políticas de saúde pública. A persistência de altas taxas dessas mortes no Brasil é "
              "considerada um indicador de iniquidade — pessoas morrem de causas que o sistema de "
              "saúde, em tese, sabe prevenir.", False, False)])

p = doc.add_paragraph()
add_rich(p, [("Hipótese. ", True, False),
             ("A taxa de mortalidade por causas evitáveis de um município está associada ao seu "
              "nível de desenvolvimento socioeconômico (renda, IDH-M, desigualdade, saneamento, "
              "escolaridade) e ao esforço público em saúde.", False, False)])

p = doc.add_paragraph()
add_rich(p, [("Objetivo. ", True, False),
             ("Construir um pipeline de dados reprodutível e aplicar técnicas de modelagem e "
              "mineração para: (1) quantificar a associação entre condições socioeconômicas e a "
              "mortalidade evitável; (2) identificar perfis típicos de município; e (3) caracterizar "
              "o padrão de causas de cada perfil.", False, False)])

# =====================================================================
# 2. FONTES DE DADOS
# =====================================================================
add_heading("2. Fontes de Dados", level=1)
add_para("As cinco fontes integradas, todas públicas:")
add_table(
    ["Fonte", "Conteúdo", "Acesso"],
    [
        ["SIM / DATASUS", "Microdados de óbitos (CID-10), por UF e ano", "FTP (.dbc)"],
        ["IBGE / Sidra", "População municipal (Censo 2022; estimativas)", "API (tab. 4714 e 6579)"],
        ["IBGE / Sidra", "PIB municipal", "API (tab. 5938)"],
        ["SICONFI / Tesouro", "Despesa municipal em saúde (função 10)", "API DCA-Anexo I-E"],
        ["IPEA / Atlas DH", "IDH-M, Gini, renda, pobreza, saneamento (Censo 2010)", "API (séries ADH_*)"],
        ["DATASUS (tabelas)", "Descrições oficiais dos códigos CID-10", "FTP (CID10.DBF)"],
    ],
    widths=[1.6, 3.4, 2.0],
)

add_para("Tratamento de lacunas temporais.", bold=True)
add_para(
    "O painel cobre 2022, 2023 e 2024, mas nem toda fonte tem cobertura completa. Adotamos "
    "substituições explícitas, registradas em uma coluna ANO_FONTE para auditoria:")
add_bullet([("População 2023: ", True, False),
            ("o IBGE não publicou estimativa; usamos o Censo 2022 como proxy.", False, False)])
add_bullet([("PIB 2024: ", True, False),
            ("indisponível por defasagem contábil (~2 anos); usamos o PIB de 2023.", False, False)])
add_bullet([("Atlas DH: ", True, False),
            ("calculado apenas com base no Censo 2010; é tratado como invariante no tempo "
             "(repetido nos três anos).", False, False)])

# =====================================================================
# 3. METODOLOGIA
# =====================================================================
add_heading("3. Metodologia", level=1)

add_heading("3.1. Classificação das causas evitáveis (LBCE)", level=2)
add_para(
    "Utilizamos a Lista Brasileira de Causas de Mortes Evitáveis de 5 a 74 anos (Malta et al., "
    "2007; atualização 2010), que organiza códigos CID-10 segundo o tipo de intervenção do SUS "
    "capaz de evitar a morte. A lista oficial possui cinco grupos; neste trabalho utilizamos quatro:")
add_table(
    ["Grupo", "Descrição"],
    [
        ["imunoprevencao", "Reduzíveis por vacinação"],
        ["doencas_infecciosas", "Tuberculose, HIV, hepatites, infecções respiratórias, doenças tropicais"],
        ["doencas_nao_transmissiveis", "Neoplasias, doenças cardiovasculares, diabetes, DPOC"],
        ["morte_materna", "Gravidez, parto e puerpério"],
    ],
    widths=[2.4, 4.6],
)
p = doc.add_paragraph()
add_rich(p, [("Exclusão das causas externas. ", True, False),
             ("O quinto grupo da lista oficial — causas externas (acidentes, quedas, agressões, "
              "suicídios) — foi deliberadamente excluído. As variáveis socioeconômicas disponíveis "
              "explicam mal essa mortalidade: acidentes de trânsito, homicídios e quedas respondem a "
              "fatores como segurança pública, infraestrutura viária e violência interpessoal, "
              "ausentes do nosso conjunto de features. Incluí-las introduziria ruído estrutural sem "
              "ganho explicativo. O escopo final é o das causas evitáveis sensíveis ao SUS e ao "
              "desenvolvimento socioeconômico.", False, False)])

add_heading("3.2. Pipeline de dados", level=2)
add_para("O projeto é organizado em notebooks sequenciais, do ETL à mineração. Cada etapa é "
         "idempotente e persiste seus resultados em disco.")
add_table(
    ["Notebook", "Função"],
    [
        ["01_extracao", "Download e cache das cinco fontes (com fallbacks de ano)."],
        ["02_transformacao", "Aplicação da LBCE, agregação e construção da feature matrix (município × ano)."],
        ["02b_transformacao_microdados", "Dataset individualizado (1 óbito por linha) com estatística descritiva."],
        ["03_regressao", "Regressão multivariada em painel (OLS, RF, XGBoost)."],
        ["04_clusterizacao", "Clusterização K-Means dos municípios."],
        ["05_mineracao", "Caracterização das causas CID-10 por cluster."],
        ["06_taxa_padronizada_idade", "Padronização etária da taxa e comparação com a bruta."],
    ],
    widths=[2.6, 4.4],
)
p = doc.add_paragraph()
add_rich(p, [("Unidade de análise e join. ", True, False),
             ("O SIM identifica o município de residência por um código de 6 dígitos (CODMUNRES); o "
              "IBGE usa 7 dígitos (com dígito verificador). A junção entre fontes é feita pelos 6 "
              "primeiros dígitos. A feature matrix final tem uma linha por município e ano "
              "(16.711 linhas; 5.570 municípios × 3 anos).", False, False)])
p = doc.add_paragraph()
add_rich(p, [("Variável-resposta: ", True, False),
             ("taxa bruta = (óbitos evitáveis de 5–74 anos ÷ população) × 100.000. Como "
              "desfecho principal, porém, usamos a taxa padronizada por idade (abaixo).", False, False)])

add_heading("3.3. Padronização por idade", level=2)
add_para(
    "A taxa bruta não corrige a estrutura etária: como as causas evitáveis (5–74 anos) são "
    "dominadas por doenças crônicas que matam perto dos 70, municípios mais envelhecidos têm "
    "taxa bruta maior só por isso. Calculamos a taxa padronizada por idade (padronização "
    "direta): a taxa específica de cada faixa quinquenal é ponderada pelos pesos da estrutura "
    "etária nacional (Censo 2022). O numerador por faixa vem dos microdados (idade de cada "
    "óbito); o denominador por faixa é extraído do Censo 2022 (IBGE, tabela 9514). É a taxa que "
    "cada município teria se tivesse a estrutura etária do Brasil — tornando as comparações "
    "entre eles justas.")

add_heading("3.4. Modelagem (regressão em painel)", level=2)
add_para(
    "Tratamos os três anos como um painel agrupado (pooled), com o ano incluído como variável "
    "dummy para capturar choques comuns de período. Optou-se por não usar efeitos fixos de "
    "município porque a maioria das features (Atlas 2010) é invariante no tempo e seria absorvida "
    "por esses efeitos. Três modelos foram comparados:")
add_bullet([("OLS ", True, False),
            ("(statsmodels), com erros-padrão cluster-robust por município — corrige a correlação "
             "entre as três observações de cada cidade.", False, False)])
add_bullet([("Random Forest e XGBoost, ", True, False),
            ("para capturar não-linearidades e interações.", False, False)])
add_para(
    "A divisão treino/teste usa GroupShuffleSplit por município (80/20), impedindo que o mesmo "
    "município apareça em treino e teste — o que inflaria artificialmente o R².")

add_heading("3.5. Clusterização", level=2)
add_para(
    "Para a tipologia de municípios, agregamos as três observações de cada município em uma única "
    "linha (média das variáveis temporais; valor único das do Atlas). Aplicamos K-Means sobre sete "
    "features socioeconômicas padronizadas (IDH-M, Gini, log do PIB per capita, log da despesa em "
    "saúde per capita, analfabetismo, água encanada e esperança de vida). O número de clusters "
    "(k=4) foi escolhido pela combinação de método do cotovelo e silhouette, priorizando a "
    "interpretabilidade. A taxa de mortalidade não entrou nas features (evita circularidade): ela "
    "é usada apenas para caracterizar os clusters depois de formados.")

add_heading("3.6. Mineração das causas", level=2)
add_para(
    "Cada óbito individual foi associado ao cluster de seu município. As descrições oficiais dos "
    "códigos CID-10 vieram das tabelas auxiliares do DATASUS (CID10.DBF). Para identificar causas "
    "características de cada cluster — e não apenas as mais frequentes — usamos o lift: a razão "
    "entre a participação da causa no cluster e sua participação nacional. Valores acima de 1 "
    "indicam sobre-representação no cluster.")

# =====================================================================
# 4. RESULTADOS
# =====================================================================
add_heading("4. Resultados", level=1)

add_heading("4.1. Panorama descritivo", level=2)
add_para(
    "Dos 4,54 milhões de óbitos registrados no triênio, 2,49 milhões estão na faixa de 5 a 74 "
    "anos, dos quais 1,37 milhão foram classificados como evitáveis pela LBCE. A distribuição "
    "entre grupos é fortemente dominada pelas doenças não transmissíveis:")
add_table(
    ["Grupo", "Óbitos", "%"],
    [
        ["Doenças não transmissíveis", "1.109.236", "81,2"],
        ["Doenças infecciosas", "250.735", "18,4"],
        ["Morte materna", "4.827", "0,4"],
        ["Imunoprevenção", "1.669", "0,1"],
        ["Total", "1.366.467", "100,0"],
    ],
    widths=[3.6, 2.0, 1.2],
)
add_para(
    "A taxa bruta média municipal foi de aproximadamente 226 óbitos evitáveis por 100 mil "
    "habitantes, com grande dispersão (de <50 a >650 por 100 mil).")
p = doc.add_paragraph()
add_rich(p, [("Efeito da padronização por idade. ", True, False),
             ("Pela taxa bruta, o gradiente regional é invertido em relação à hipótese "
              "(Sul/Sudeste no topo). A padronização corrige o paradoxo: o Nordeste passa a "
              "superar Sudeste e Sul, e o Norte sobe de 159 para 228 — coerente com a tese.", False, False)])
add_table(
    ["Região", "Taxa bruta", "Taxa padronizada"],
    [
        ["Norte", "158,9", "227,6"],
        ["Nordeste", "204,1", "247,3"],
        ["Centro-Oeste", "229,8", "257,4"],
        ["Sudeste", "244,6", "244,5"],
        ["Sul", "256,5", "241,9"],
    ],
    widths=[2.2, 2.0, 2.4],
)

add_heading("4.2. Regressão", level=2)
add_para(
    "Os modelos foram estimados sobre a taxa padronizada (desfecho principal). O desempenho "
    "preditivo é baixo (R² ≈ 0,09–0,10): ao remover a variância de estrutura etária — que estava "
    "correlacionada com região/desenvolvimento e inflava o R² da taxa bruta (≈ 0,23) — resta "
    "apenas o sinal socioeconômico puro, menor mas no sentido correto. O objetivo é inferência, "
    "não predição.")
add_table(
    ["Modelo", "R²", "RMSE", "MAE"],
    [
        ["OLS (linear)", "0,092", "71,0", "53,9"],
        ["Random Forest", "0,091", "71,0", "54,1"],
        ["XGBoost", "0,103", "70,6", "53,7"],
    ],
    widths=[2.6, 1.4, 1.4, 1.4],
)
add_para(
    "No painel completo (com controles de região e ano), o Gini permanece o achado central: "
    "+151 (p<0,001) — mais desigualdade, mais mortalidade. Renda per capita (−0,14), esperança "
    "de vida (−4,2) e proporção de pobres (−2,1, colinear com renda) também são significativas. "
    "O efeito do IDH-M é absorvido pelas dummies de região. Para isolá-lo, a tabela abaixo "
    "compara um modelo transversal sem controles de região, com alvo bruto vs. padronizado "
    "(*** p<0,001):")
add_table(
    ["Variável", "Alvo: taxa bruta", "Alvo: taxa padronizada"],
    [
        ["IDH-M", "−77,1 ***", "−126,5 ***"],
        ["Gini", "+43,6 *", "+85,8 ***"],
        ["Esperança de vida", "−0,1", "−3,5 ***"],
    ],
    widths=[2.4, 2.2, 2.4],
)
add_para(
    "Com a taxa padronizada, o IDH-M passa a ter efeito negativo e fortemente significativo "
    "(mais desenvolvimento → menos mortalidade) e o Gini reforça — exatamente os sinais previstos "
    "pela hipótese. A importância das features nos modelos de árvore é liderada pela proporção de "
    "pobres e pela região.")

add_heading("4.3. Clusters de municípios", level=2)
add_para("A clusterização produziu quatro perfis bem separados (ANOVA altamente significativa), "
         "renumerados por IDH-M crescente (cluster 0 = mais vulnerável). A taxa padronizada é a "
         "comparável:")
add_table(
    ["Cl.", "Mun.", "Taxa bruta", "Taxa padr.", "IDH-M", "Gini", "Água%", "Região dom."],
    [
        ["0", "784", "183", "232", "0,56", "0,54", "58", "NE/N"],
        ["1", "1.573", "208", "249", "0,60", "0,52", "84", "NE"],
        ["2", "1.430", "248", "240", "0,70", "0,45", "92", "SE/S"],
        ["3", "1.774", "243", "251", "0,72", "0,48", "94", "SE/S"],
    ],
    widths=[0.5, 0.85, 1.0, 1.0, 0.85, 0.75, 0.75, 1.1],
)
add_para(
    "Os clusters 0 e 1 reúnem municípios de menor IDH-M e pior saneamento (Nordeste e Norte); "
    "os 2 e 3, os mais desenvolvidos (Sudeste e Sul). Pela taxa bruta, os desenvolvidos pareciam "
    "ter mortalidade maior (artefato etário); pela padronizada, as taxas se aproximam (232–251) e "
    "a falsa vantagem das regiões pobres desaparece.")

add_heading("4.4. Mineração: do que morre cada perfil", level=2)
add_para(
    "A composição agregada de causas é semelhante entre clusters — todos dominados por infarto, "
    "diabetes e neoplasias. O sinal socioeconômico aparece nas causas sobre-representadas (lift):")
add_table(
    ["Cluster", "Causas características (lift)"],
    [
        ["0 (mais pobre)", "AVC (1,6×), transtornos por álcool (1,6×), diabetes, doença de Chagas, diarreia"],
        ["1", "esquistossomose (2,5×), doença de Chagas (1,7×), álcool, infecções intestinais"],
        ["2 (mais rico)", "cardiopatia isquêmica crônica, infarto cerebral, neoplasia de cólon, hepatite crônica"],
        ["3", "dengue, enfisema, hipertensão, neoplasias de esôfago e laringe"],
    ],
    widths=[1.5, 5.5],
)
add_para(
    "Há, portanto, uma assinatura socioeconômica dupla: perfis pobres ainda carregam doenças "
    "infecciosas e negligenciadas (Chagas, esquistossomose, diarreia) que o desenvolvimento "
    "deveria ter eliminado, enquanto perfis ricos concentram doenças crônicas. A morte materna "
    "também segue o gradiente esperado, caindo de 0,65% dos óbitos no cluster 0 para 0,31% no "
    "cluster 3.")

# =====================================================================
# 5. DISCUSSÃO E LIMITAÇÕES
# =====================================================================
add_heading("5. Discussão e Limitações", level=1)
p = doc.add_paragraph()
add_rich(p, [("Padronização por idade (resolvida). ", True, False),
             ("A taxa bruta produzia um resultado paradoxal — maior nos clusters desenvolvidos — "
              "por confundimento etário: municípios desenvolvidos têm população mais envelhecida na "
              "faixa 5–74 anos e acumulam mais mortes por doenças crônicas. A padronização direta "
              "(com a estrutura etária do Censo 2022) corrige esse viés: o gradiente regional se "
              "realinha (Nordeste passa a superar Sudeste/Sul) e o coeficiente do IDH-M se torna "
              "negativo, como previsto. Permanecem duas ressalvas técnicas: (i) a padronização "
              "direta é sensível a ruído em municípios pequenos com poucas mortes por faixa — uma "
              "padronização indireta (SMR) seria mais estável; (ii) o denominador por idade é do "
              "Censo 2022, aplicado aos três anos (a estrutura etária varia pouco no período).", False, False)])
p = doc.add_paragraph()
add_rich(p, [("Subnotificação em áreas remotas. ", True, False),
             ("O cluster mais vulnerável (0, Norte/Nordeste rural) apresenta a menor taxa "
              "padronizada, o que provavelmente reflete subnotificação de óbitos e causas "
              "mal-definidas nessas áreas, e não menor mortalidade real — uma limitação dos "
              "próprios registros do SIM.", False, False)])
p = doc.add_paragraph()
add_rich(p, [("Endogeneidade da despesa em saúde. ", True, False),
             ("O coeficiente positivo da despesa per capita (mais gasto associado a mais mortes) "
              "não deve ser lido como causal: municípios com maior carga de doença tendem a gastar "
              "mais (causalidade reversa). Isolar esse efeito exigiria instrumentos ou defasagens "
              "temporais.", False, False)])
p = doc.add_paragraph()
add_rich(p, [("Atlas DH defasado. ", True, False),
             ("IDH-M, Gini e renda referem-se a 2010, enquanto os óbitos são de 2022–2024. Essa "
              "defasagem de medição atenua as associações e contribui para o R² modesto.", False, False)])
p = doc.add_paragraph()
add_rich(p, [("Natureza ecológica. ", True, False),
             ("Os clusters e as associações são de nível municipal. Não se pode inferir, a partir "
              "deles, o comportamento de indivíduos (falácia ecológica).", False, False)])

# =====================================================================
# 6. CONCLUSÃO
# =====================================================================
add_heading("6. Conclusão", level=1)
add_para(
    "O projeto entregou um pipeline reprodutível que integra cinco fontes públicas e sustenta três "
    "análises complementares. As evidências apontam, de forma consistente, que a mortalidade por "
    "causas evitáveis no Brasil tem natureza socioeconômica: a desigualdade (Gini) está associada "
    "positivamente à taxa de forma robusta, e o tipo de causa evitável varia sistematicamente com "
    "o perfil do município — doenças infecciosas e negligenciadas nos territórios pobres, doenças "
    "crônicas nos desenvolvidos. A padronização por idade, incorporada ao pipeline, corrigiu o "
    "paradoxo das taxas brutas e alinhou os sinais da regressão à hipótese (IDH-M negativo, Gini "
    "positivo). O aprofundamento futuro mais relevante é a incorporação de variáveis da rede de "
    "atenção à saúde (leitos e médicos por habitante, cobertura da atenção básica) e o tratamento "
    "da subnotificação em áreas remotas.")

# =====================================================================
# REFERÊNCIAS
# =====================================================================
add_heading("Referências", level=1)
for ref in [
    "Malta DC et al. Lista de causas de mortes evitáveis por intervenções do SUS. "
    "Epidemiol. Serv. Saúde, 16(4):233–244, 2007.",
    "Malta DC et al. Atualização da lista de causas de mortes evitáveis [Nota técnica]. "
    "Epidemiol. Serv. Saúde, 19(2):173–176, 2010.",
    "DATASUS. Sistema de Informações sobre Mortalidade (SIM) — microdados CID-10.",
    "IBGE. Sidra — População e PIB municipal.",
    "Tesouro Nacional. SICONFI — Declaração de Contas Anuais (DCA).",
    "IPEA. Atlas do Desenvolvimento Humano (Censo 2010).",
]:
    doc.add_paragraph(ref, style="List Bullet")

doc.save(OUT)
print(f"Salvo: {OUT}  ({OUT.stat().st_size/1024:.1f} KB)")
